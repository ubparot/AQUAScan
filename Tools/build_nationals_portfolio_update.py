from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "aquascan-nationals-update"
DOCX_PATH = OUT / "AquaScan_Nationals_Update_Pages.docx"

WEB_PLAN = ROOT / "web" / "tmp-project-preflight.png"
WEB_AI = ROOT / "web" / "tmp-research-analysis.png"
WEB_TELEMETRY = ROOT / "web" / "tmp-telemetry-ui.png"
BRIDGE = ROOT / ".codex_bridge_now.png"

BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4F9"
GRAY = "E7E6E6"
GREEN = "E2F0D9"
ORANGE = "FCE4D6"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=9.5, bold=False, color="000000", italic=False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, after=3, before=0, line=1.05, keep=False) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep


def add_page_number(section, label: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(label)
    set_font(run, size=9)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=8, keep=True)
    run = p.add_run(text)
    set_font(run, size=13, bold=True)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=5, after=3, keep=True)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=True)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=5)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)


def add_bullets(doc: Document, items: list[tuple[str, str] | str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        format_paragraph(p, after=2)
        if isinstance(item, tuple):
            lead, text = item
            r1 = p.add_run(lead)
            set_font(r1, bold=True)
            r2 = p.add_run(text)
            set_font(r2)
        else:
            run = p.add_run(item)
            set_font(run)


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(7.35)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    format_paragraph(p, after=0)
    r1 = p.add_run(label)
    set_font(r1, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], small=False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, after=0)
        run = p.add_run(text)
        set_font(run, size=8 if small else 8.5, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Inches(widths[index])
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else PALE_BLUE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            format_paragraph(p, after=0, line=1.0)
            run = p.add_run(value)
            set_font(run, size=7.5 if small else 8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, image: Path, caption: str, width=7.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, after=2)
    p.add_run().add_picture(str(image), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(c, after=5)
    run = c.add_run(caption)
    set_font(run, size=7.5, bold=True)


def new_page(doc: Document, label: str) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    add_page_number(section, label)


def configure_section(section) -> None:
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)


def build_docx() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc.sections[0])
    add_page_number(doc.sections[0], "1A")

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    add_title(doc, "NATIONALS DEVELOPMENT UPDATE MAP")
    add_body(
        doc,
        "AquaScan did not remain frozen after the TSA States submission. The earlier portfolio documents the system that proved the core concept: a mobile vessel, deployable probe, Unity visualization, and field data workflow. The pages identified below document the next engineering cycle completed for Nationals.",
    )
    add_callout(
        doc,
        "How to read this revision: ",
        "the original pages remain as evidence of the States baseline. These inserted pages show the problem identified after States, the refinement made, and the verified result.",
    )
    add_heading(doc, "Post-States Development Sequence")
    add_table(
        doc,
        ["Stage", "Engineering Focus", "Result"],
        [
            ["States baseline", "Prove collection, control, and visualization concept", "Unity-based interface and working depth-aware data workflow"],
            ["Field-access refinement", "Reduce setup friction and improve operator access", "Browser-based mission-control dashboard"],
            ["Control-chain refinement", "Improve communication reliability and safety feedback", "WebSocket, ESP32, Arduino Mega, ESC, winch, and RS-485 telemetry chain"],
            ["Data-use refinement", "Move from viewing measurements toward interpreting patterns", "MARIS experimental predictive-analysis model and research-analysis tools"],
            ["Verification", "Test software behavior and model pipeline", "26 web tests, production build, lint, and 7 ML tests passed on June 14, 2026"],
        ],
        [1.15, 3.0, 3.2],
    )
    add_heading(doc, "Inserted Page Guide")
    add_bullets(
        doc,
        [
            ("Pages 16A-16B: ", "web mission control and embedded communications iteration."),
            ("Page 16C: ", "MARIS experimental predictive analysis, placed directly before “What This Data Can Reveal.”"),
            ("Page 28A: ", "problem-change-evidence-impact iteration record and current verification results."),
            ("Page 29A: ", "post-States development record showing continued engineering activity."),
        ],
    )
    add_heading(doc, "Current System Position")
    add_body(
        doc,
        "AquaScan is currently a working integrated prototype for calm, shallow, and sheltered water environments. Manual control, telemetry, mission replay, data visualization, probe communications, and the training/export path for MARIS have been implemented. Full autonomous route execution and field-validated MARIS inference remain future development.",
    )

    new_page(doc, "16A")
    add_title(doc, "POST-STATES ITERATION: FROM UNITY PROTOTYPE TO WEB MISSION CONTROL")
    add_body(
        doc,
        "The Unity interface shown on the preceding original portfolio page was the States baseline. It demonstrated that AquaScan data could be replayed, mapped, and connected to the vessel. After States, the team identified a field-use limitation: Unity required a dedicated installed application and made it harder to open the controls quickly on different operator devices.",
    )
    add_callout(
        doc,
        "Iteration decision: ",
        "retain the proven mission-data and control concepts, but rebuild the operator interface as a browser-based field dashboard that can run locally and be opened from devices on the same network.",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["States Baseline", "Post-States Web Refinement", "Engineering Benefit"],
        [
            ["Unity desktop application", "Local browser-based dashboard", "Faster access and no Unity installation on the operator device"],
            ["Mission replay and map layers", "Replay, route planning, saved project files, and exportable bundles", "Supports repeatable mission preparation and review"],
            ["Live control concept", "Direct WebSocket connection to the boat with visible connection state", "Improves field awareness and troubleshooting"],
            ["Single presentation-focused layout", "Simple/advanced drive modes, light/dark themes, and tabbed controls", "Separates essential control from detailed diagnostics"],
        ],
        [2.15, 2.7, 2.5],
        small=True,
    )
    add_image(doc, WEB_PLAN, "Current web mission planner showing saved route data, validation status, playback, and telemetry.")
    add_bullets(
        doc,
        [
            ("Preserved capability: ", "CSV/JSON mission loading, GPS route projection, sample points, heat maps, and timeline playback."),
            ("Added capability: ", "mission-plan editing, local project saving, preflight checks, direct live control, and browser-accessible operation."),
            ("Current validation: ", "the web application passes 26 automated tests, production build, and lint checks."),
        ],
    )

    new_page(doc, "16B")
    add_title(doc, "WEB FIELD SYSTEM AND EMBEDDED COMMUNICATIONS ITERATION")
    add_body(
        doc,
        "The interface change was paired with a deeper electronics and firmware refinement. The post-States system separates high-level operator commands, motor control, and probe sensing across communication links chosen for each task.",
    )
    add_heading(doc, "Current End-to-End Control and Data Path")
    add_table(
        doc,
        ["Link", "Current Function", "Reason for Refinement"],
        [
            ["Browser dashboard → ESP32", "WebSocket commands, status, arming, E-stop, and live telemetry", "Provides direct network control and immediate operator feedback"],
            ["ESP32 → Arduino Mega", "Full-duplex hardware UART control bridge", "More reliable than the earlier bridge approach and keeps USB debug available"],
            ["Arduino Mega → ESCs/winch", "Applies motor pulses and deployment commands", "Centralizes real-time actuator control and neutral-output behavior"],
            ["Arduino Mega ↔ probe", "Half-duplex RS-485 sensor communication", "Supports longer tether communication and separates sensor traffic from drive control"],
            ["Telemetry → dashboard", "Sensor values, battery, RSSI, depth, motor output, and status", "Turns hidden system state into visible diagnostic evidence"],
        ],
        [1.55, 3.0, 2.8],
        small=True,
    )
    add_image(doc, WEB_TELEMETRY, "Browser mission-control dashboard showing route visualization, operator controls, and live system status.")
    add_heading(doc, "Safety and Reliability Behaviors")
    add_bullets(
        doc,
        [
            ("Safe startup: ", "the system begins disarmed and commands neutral motor output."),
            ("Loss-of-command response: ", "disconnects, malformed messages, or command timeout force neutral output."),
            ("Operator override: ", "E-stop is visible in the dashboard and latches until reset."),
            ("Debug visibility: ", "Arduino and ESP32 status output makes command sequence, arm state, winch state, and motor outputs inspectable."),
        ],
    )

    new_page(doc, "16C")
    add_title(doc, "MARIS: EXPERIMENTAL PREDICTIVE ANALYSIS")
    add_body(
        doc,
        "MARIS is AquaScan’s first-pass machine-learning model for turning structured mission data into predictive research signals. It is an implemented proof of concept, not a field-validated decision system. MARIS demonstrates how future AquaScan missions could move beyond displaying measurements and begin forecasting or flagging conditions that deserve closer investigation.",
    )
    add_callout(
        doc,
        "Accurate framing: ",
        "MARIS has been trained, evaluated, and exported as an ONNX model. The current primary dashboard still uses a transparent heuristic fallback while larger independent field datasets and live model integration are developed.",
        fill=ORANGE,
    )
    add_heading(doc, "What MARIS Produces")
    add_bullets(
        doc,
        [
            ("Current dissolved oxygen estimate: ", "provides the model’s present-condition regression output."),
            ("Dissolved oxygen forecasts: ", "predicts values at +30, +60, and +120 minutes."),
            ("Bloom-risk probability: ", "identifies combinations of conditions associated with elevated bloom risk."),
            ("Anomaly probability: ", "flags unusual or potentially concerning water-quality patterns."),
        ],
    )
    add_table(
        doc,
        ["Development Evidence", "Current Result", "Interpretation"],
        [
            ["Training data", "3 real missions + 12 synthetic missions; 1,158 windows", "Sufficient for proof of concept, not broad field validation"],
            ["Evaluation data", "232 evaluation windows", "Preliminary internal evaluation"],
            ["Oxygen RMSE", "1.92 mg/L", "Shows the current model still requires accuracy improvement"],
            ["Forecast RMSE", "2.08 mg/L", "Useful as a development benchmark, not a validated forecast claim"],
            ["Bloom accuracy", "59.5%", "Early-stage result; larger labeled datasets are required"],
            ["Anomaly accuracy", "93.5%", "Promising internal result, subject to label and dataset limitations"],
            ["ONNX export", "Validated; max output difference 0.00000095", "Confirms exported model matches the PyTorch output"],
        ],
        [1.65, 2.5, 3.2],
        small=True,
    )
    add_image(doc, WEB_AI, "Current dashboard research-analysis interface; the visible fallback label prevents experimental output from being presented as validated inference.")

    new_page(doc, "28A")
    add_title(doc, "NATIONALS ITERATION EVIDENCE: PROBLEM → CHANGE → EVIDENCE → IMPACT")
    add_body(
        doc,
        "The largest post-States refinement was not one isolated feature. It was a system-level iteration focused on field usability, communications reliability, diagnostic visibility, and future data interpretation.",
    )
    add_table(
        doc,
        ["Problem Identified", "Refinement Made", "Evidence Available", "Impact"],
        [
            ["Unity was effective for demonstration but less convenient across field devices.", "Rebuilt the operator interface as a local React/Vite web dashboard.", "Current dashboard screenshots; passing build and automated tests.", "Faster access and a clearer path to use on laptops or tablets."],
            ["The earlier communications chain exposed limited diagnostic information.", "Added explicit WebSocket status, sequence tracking, hardware UART bridge, and RS-485 probe telemetry.", "Firmware source, serial-monitor output, and UI telemetry cards.", "Makes failures easier to locate and system state easier to verify."],
            ["Operator controls needed stronger safety behavior.", "Added disarmed startup, neutral timeout behavior, arm/disarm state, and latched E-stop.", "Control contract and automated drive-mapping tests.", "Reduces risk of unintended motor output during setup or communication loss."],
            ["Mission data existed as CSV/JSON but was difficult to organize as a repeatable project.", "Added mission planning, validation, local project saving, and exportable bundles.", "Saved-plan UI and route-tool tests.", "Supports repeatable setup, review, and communication of mission data."],
            ["Visualizations explained the past but did not explore predictive use.", "Developed MARIS training/export pipeline and research-analysis scaffolding.", "PyTorch/ONNX artifacts, metrics report, and 7 passing ML tests.", "Demonstrates a realistic future path while clearly documenting present limitations."],
        ],
        [1.6, 2.05, 1.85, 1.85],
        small=True,
    )
    add_heading(doc, "Verified June 14, 2026")
    add_bullets(
        doc,
        [
            ("Web tests: ", "26 passed across mission loading, route planning, project files, control math, research metadata, and RS-485 sensor averaging."),
            ("Web release checks: ", "production build completed and lint completed without errors."),
            ("ML checks: ", "7 tests passed across data loading, feature engineering, normalization, and training outputs."),
            ("Model export: ", "ONNX artifact exported and numerically validated against the PyTorch model."),
        ],
    )
    add_callout(
        doc,
        "Engineering significance: ",
        "the system now preserves the functioning States prototype as evidence while showing a traceable cycle of limitation identification, redesign, implementation, and verification.",
        fill=GREEN,
    )

    new_page(doc, "29A")
    add_title(doc, "POST-STATES DEVELOPMENT RECORD")
    add_body(
        doc,
        "The following record supplements the original TSA work log by documenting major repository-backed development completed after the April 29, 2026 States portfolio submission.",
    )
    add_table(
        doc,
        ["Date", "Development Activity", "Result / Evidence"],
        [
            ["04/30/26", "Expanded live-control WebSocket and metrics integration.", "Control and data behavior incorporated into the evolving software system."],
            ["05/05/26", "Developed and tested first-pass multi-task ML pipeline.", "MARIS training code, PyTorch checkpoint, ONNX export, normalization data, and metrics report."],
            ["05/06/26", "Created the browser-based AquaScan dashboard.", "React/Vite/Three.js dashboard with mission loading, visualization, drive controls, and tests."],
            ["05/08-05/10/26", "Integrated project-wide software, visualization, and control updates.", "Consolidated system capable of playback, live control, and expanded data interpretation."],
            ["05/28/26", "Expanded planning, project-file, preflight, and research-analysis tools.", "Saved routes, validation feedback, project bundles, and research-analysis scaffolding."],
            ["06/07/26", "Added RS-485 sensor pipeline and UI telemetry.", "Separated probe sensing from motor-control communications and surfaced telemetry in the dashboard."],
            ["06/11-06/12/26", "Validated current bridge/control behavior and captured evidence.", "Serial-monitor and connected-dashboard evidence of the control chain."],
            ["06/14/26", "Ran current verification suite and prepared Nationals portfolio update.", "26 web tests passed; build/lint passed; 7 ML tests passed."],
        ],
        [1.05, 3.35, 3.0],
        small=True,
    )
    add_heading(doc, "Current Limitations Carried Forward")
    add_bullets(
        doc,
        [
            ("MARIS validation: ", "requires substantially more independent real-water mission data and external validation."),
            ("Autonomy: ", "manual control and mission planning are implemented; full autonomous route execution remains a planned refinement."),
            ("Depth accuracy: ", "probe depth is still estimated from tether deployment rather than measured with a pressure sensor."),
            ("Sensor accuracy: ", "continued calibration and optical-window refinement are required for stronger scientific claims."),
        ],
    )
    add_body(
        doc,
        "These limitations are included because the purpose of iteration is not to claim that every problem is solved. They define the next measurable engineering cycle for AquaScan.",
        bold_lead="These limitations",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
